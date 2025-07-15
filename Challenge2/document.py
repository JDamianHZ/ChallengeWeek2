from abc import ABC, abstractmethod
from docx import Document as DocxDocument  # para Word
from fpdf import FPDF                      # para PDF

# Clase abstracta base
class Document(ABC):
    @abstractmethod
    def read(self) -> str:
        pass

    @abstractmethod
    def save(self, filename: str):
        pass

# PDFDocument - crea archivo PDF real
class PDFDocument(Document):
    def read(self) -> str:
        return "Reading PDF Document..."

    def save(self, filename: str = "document.pdf"):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="This is a sample PDF document.", ln=True)
        pdf.output(filename)
        print(f"PDF saved as {filename}")

# WordDocument - crea archivo DOCX real
class WordDocument(Document):
    def read(self) -> str:
        return "Reading Word Document..."

    def save(self, filename: str = "document.docx"):
        doc = DocxDocument()
        doc.add_paragraph("This is a sample Word document.")
        doc.save(filename)
        print(f"Word document saved as {filename}")

# TextDocument - crea archivo TXT real
class TextDocument(Document):
    def read(self) -> str:
        return "Reading Text Document..."

    def save(self, filename: str = "document.txt"):
        with open(filename, "w") as f:
            f.write("This is a sample text document.")
        print(f"Text document saved as {filename}")
